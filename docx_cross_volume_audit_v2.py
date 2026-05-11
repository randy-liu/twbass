import re
import unicodedata
from collections import OrderedDict
from pathlib import Path

from docx import Document

BASE = Path(r"C:\Research\twbass-data")
FILES = OrderedDict([
    ("0A1", BASE / "0A1_台灣氣候forcing與區域差異研究.docx"),
    ("0B", BASE / "0B_南北成土母質與地球化學基底.docx"),
    ("0C1", BASE / "0C１_六大水體 seasonal 評估.docx"),
    ("0D1", BASE / "0D1_基底資料矩陣與極端事件整合.docx"),
    ("1A", BASE / "1A_短時間環境觸發與生理限制.docx"),
    ("1B1", BASE / "1B1_六大水域棲位模型與風生流.docx"),
    ("2A", BASE / "2A_覓食偏好、印記與反射咬餌.docx"),
    ("2B", BASE / "2B_側線、內耳與水下聲學傳遞.docx"),
    ("2C", BASE / "2C_視線軸向、攻擊角度與假餌操作.docx"),
    ("3A", BASE / "3A_高壓舊魚心理機制與誘咬本質.docx"),
    ("3B1", BASE / "3B1_極端情境高壓策略推演.docx"),
    ("4A", BASE / "4A_繁衍地球化學與水文干擾.docx"),
    ("4B", BASE / "4B_棲位競爭、容載量與護巢防禦.docx"),
    ("SUP-A", BASE / "SUP-A：感官生理閾值補充研究.docx"),
    ("SUP-B", BASE / "SUP-B：底棲水化學梯度補充研究.docx"),
])

TRANSLATION = str.maketrans({
    "–": "-",
    "—": "-",
    "−": "-",
    "〜": "~",
    "～": "~",
    "：": ":",
    "（": "(",
    "）": ")",
    "，": ",",
    "；": ";",
    "％": "%",
    "　": " ",
})
FLAGS = re.I | re.S
CODE_RE = re.compile(r"\b(?:B0-\d{2}|V[1-4][ABC]-\d{2}|VSUP-[AB]-?\d{2}|CF-\d{2})\b")
NUM_RE = re.compile(r"(?:>=|<=|>|<|~|≈)?\s*\d+(?:\.\d+)?(?:\s*[-~]\s*\d+(?:\.\d+)?)?(?:\s*(?:cm|m|mg/L|mV|Hz|ng/mL|ng/ml|hr|h|min|day|天|月下旬|°C|C))?", re.I)


def norm(text: str) -> str:
    text = unicodedata.normalize("NFKC", text or "")
    text = text.translate(TRANSLATION)
    text = text.replace("\xa0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\s*\n\s*", " ", text)
    return text.strip()


def read_doc(path: Path):
    doc = Document(path)
    lines = []
    for para in doc.paragraphs:
        t = norm(para.text)
        if t:
            lines.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = norm(cell.text)
                if t:
                    lines.append(t)
    deduped = []
    prev = None
    for line in lines:
        if line != prev:
            deduped.append(line)
        prev = line
    return {"lines": deduped, "text": "\n".join(deduped)}


def re_has(text, pattern):
    return bool(re.search(pattern, text, FLAGS))


def re_any(text, patterns):
    return any(re_has(text, p) for p in patterns)


def match_all_groups(text, groups):
    return all(re_any(text, group) for group in groups)


def find_snippets_regex(lines, patterns, radius=1, limit=5):
    out = []
    seen = set()
    for idx, line in enumerate(lines):
        if any(re.search(p, line, FLAGS) for p in patterns):
            start = max(0, idx - radius)
            end = min(len(lines), idx + radius + 1)
            snippet = " | ".join(lines[start:end])
            snippet = re.sub(r"\s+", " ", snippet).strip()
            if snippet not in seen:
                seen.add(snippet)
                out.append(snippet)
            if len(out) >= limit:
                break
    return out


def first_matches(text, patterns):
    found = []
    for label, pat in patterns:
        if re_has(text, pat):
            found.append(label)
    return found


def generic_values(snippets):
    values = []
    for snippet in snippets:
        for m in NUM_RE.findall(snippet):
            m = m.strip()
            if not m or re.fullmatch(r"\d+", m):
                continue
            values.append(m)
    uniq = []
    seen = set()
    for v in values:
        if v not in seen:
            seen.add(v)
            uniq.append(v)
    return uniq[:8]


def classify_parameter(doc_key, info, spec):
    text = info["text"]
    lines = info["lines"]
    expected_here = doc_key in spec.get("expected_docs", [])
    scan_docs = spec.get("scan_docs") or spec.get("expected_docs", [])
    if doc_key not in scan_docs and not spec.get("scan_all"):
        return None
    labels = first_matches(text, spec["checks"])
    anchor_groups = spec.get("anchor_groups", [])
    flat_anchor_patterns = [p for group in anchor_groups for p in group]
    anchors_present = match_all_groups(text, anchor_groups) if anchor_groups else bool(labels)
    snippets = find_snippets_regex(lines, flat_anchor_patterns or [pat for _, pat in spec["checks"]], radius=1, limit=5)
    if labels:
        status = "符合"
        if spec.get("mode") == "all" and len(labels) < len(spec["checks"]):
            status = "偏差"
        return {"value": "; ".join(labels), "status": status, "snippets": snippets}
    if anchors_present:
        vals = generic_values(snippets)
        value = "; ".join(vals) if vals else "有提及但未抽出精確值"
        return {"value": value, "status": "偏差", "snippets": snippets}
    if expected_here:
        return {"value": "未找到", "status": "缺失", "snippets": []}
    return None


def section_slice(lines, start_markers, end_markers):
    start = None
    for i, line in enumerate(lines):
        if any(marker.lower() in line.lower() for marker in start_markers):
            start = i
            break
    if start is None:
        return []
    end = len(lines)
    for i in range(start + 1, len(lines)):
        if any(marker.lower() in lines[i].lower() for marker in end_markers):
            end = i
            break
    return lines[start:end]


def canonical_code(code: str) -> str:
    code = norm(code).upper().replace(" ", "")
    code = code.replace("VSUP-A-", "VSUP-A").replace("VSUP-B-", "VSUP-B")
    return code


def code_target(code: str):
    code = canonical_code(code)
    mapping = {
        "B0": "0D1",
        "V1A": "1A",
        "V1B": "1B1",
        "V2A": "2A",
        "V2B": "2B",
        "V2C": "2C",
        "V3A": "3A",
        "V3B": "3B1",
        "V4A": "4A",
        "V4B": "4B",
        "VSUP-A": "SUP-A",
        "VSUP-B": "SUP-B",
    }
    for prefix, target in mapping.items():
        if code.startswith(prefix):
            return target
    return None


def code_exists_in_doc(code: str, text: str):
    code = canonical_code(code)
    patterns = [re.escape(code)]
    if code.startswith("VSUP-A"):
        patterns.append(re.escape(code.replace("VSUP-A", "VSUP-A-")))
    if code.startswith("VSUP-B"):
        patterns.append(re.escape(code.replace("VSUP-B", "VSUP-B-")))
    return any(re_has(text, p) for p in patterns)


def snippets_for_code(code: str, lines):
    code = canonical_code(code)
    pats = [re.escape(code)]
    if code.startswith("VSUP-A"):
        pats.append(re.escape(code.replace("VSUP-A", "VSUP-A-")))
    if code.startswith("VSUP-B"):
        pats.append(re.escape(code.replace("VSUP-B", "VSUP-B-")))
    return find_snippets_regex(lines, pats, radius=1, limit=3)


def normalize_code_set(codes):
    out = []
    seen = set()
    for c in codes:
        cc = canonical_code(c)
        if cc not in seen:
            seen.add(cc)
            out.append(cc)
    return out


def has_any(text: str, needles):
    low = text.lower()
    return any(n.lower() in low for n in needles)


def find_plain_snippets(lines, needles, radius=1, limit=5):
    pats = [re.escape(n) for n in needles]
    return find_snippets_regex(lines, pats, radius=radius, limit=limit)


def main():
    docs = {k: read_doc(v) for k, v in FILES.items()}

    param_specs = [
        {
            "name": "ART(Alert Reset Time)常態",
            "expected": "22°C=20-24 hr; 15°C=36-42 hr; 30°C=12 hr",
            "expected_docs": ["3A", "SUP-A"],
            "scan_docs": ["3A", "SUP-A", "3B1"],
            "anchor_groups": [[r"Alert Reset(?: Time)?", r"\bART\b", r"警覺重置"]],
            "mode": "all",
            "checks": [
                ("22°C=20-24 hr", r"22\s*°?\s*C.{0,120}?20\s*[-~]\s*24\s*(?:h|hr)"),
                ("15°C=36-42 hr", r"15\s*°?\s*C.{0,120}?36\s*[-~]\s*42\s*(?:h|hr)"),
                ("30°C=12 hr", r"30\s*°?\s*C.{0,120}?12\s*(?:h|hr)")
            ],
        },
        {
            "name": "ART H2S縮短版",
            "expected": "80-120 min",
            "expected_docs": ["3A"],
            "scan_docs": ["3A", "3B1", "SUP-A"],
            "anchor_groups": [[r"Alert Reset(?: Time)?", r"\bART\b", r"警覺重置"], [r"H2S", r"硫化氫"]],
            "checks": [("80-120 min", r"80\s*[-~]\s*120\s*min")],
        },
        {
            "name": "皮質醇LVF基線",
            "expected": ">20-40 ng/mL",
            "expected_docs": ["SUP-A", "3A"],
            "scan_docs": ["SUP-A", "3A", "3B1"],
            "anchor_groups": [[r"LVF", r"Low Vulnerability"], [r"皮質醇", r"cortisol"]],
            "checks": [(">20-40 ng/mL", r">?\s*20\s*[-~]\s*40\s*ng\s*/\s*mL")],
        },
        {
            "name": "CFF明視覺22°C",
            "expected": "30-60 Hz",
            "expected_docs": ["2B", "SUP-A"],
            "scan_docs": ["2B", "SUP-A"],
            "anchor_groups": [[r"\bCFF\b", r"Critical Flicker"], [r"明視", r"photopic"]],
            "checks": [("30-60 Hz", r"30\s*[-~]\s*60\s*Hz")],
        },
        {
            "name": "CFF明視覺35°C",
            "expected": "38-45 Hz",
            "expected_docs": ["SUP-A"],
            "scan_docs": ["SUP-A"],
            "anchor_groups": [[r"\bCFF\b", r"Critical Flicker"], [r"35\s*°?\s*C"], [r"明視", r"photopic"]],
            "checks": [("38-45 Hz", r"38\s*[-~]\s*45\s*Hz")],
        },
        {
            "name": "CFF暗視覺",
            "expected": "5-15 Hz",
            "expected_docs": ["SUP-A"],
            "scan_docs": ["SUP-A"],
            "anchor_groups": [[r"\bCFF\b", r"Critical Flicker"], [r"暗視", r"scotopic"]],
            "checks": [("5-15 Hz", r"5\s*[-~]\s*15\s*Hz")],
        },
        {
            "name": "SNs峰值",
            "expected": "~20 Hz(<30 Hz)",
            "expected_docs": ["2B"],
            "scan_docs": ["2B", "4B"],
            "anchor_groups": [[r"\bSNs\b", r"表淺神經丘", r"Superficial neuromasts"]],
            "checks": [("~20 Hz(<30 Hz)", r"(?:~|≈)?\s*20\s*Hz.{0,120}?<\s*30\s*Hz|<\s*30\s*Hz.{0,120}?(?:~|≈)?\s*20\s*Hz")],
        },
        {
            "name": "Follower Rejection近點",
            "expected": "13.5-24 cm",
            "expected_docs": ["2C", "3A"],
            "scan_docs": ["2C", "3A", "3B1"],
            "anchor_groups": [[r"Follower Rejection", r"Near point", r"近點"]],
            "checks": [("13.5-24 cm", r"13\.5\s*[-~]\s*24\s*cm")],
        },
        {
            "name": "Schreckstoff半徑 Zone-A/B",
            "expected": "<0.5 m",
            "expected_docs": ["3A", "SUP-B"],
            "scan_docs": ["3A", "SUP-B", "3B1"],
            "anchor_groups": [[r"Schreckstoff", r"警報物質"], [r"Zone[- ]?A", r"Zone[- ]?B", r"A/B"]],
            "checks": [("<0.5 m", r"<\s*0\.5\s*m")],
        },
        {
            "name": "Schreckstoff半徑 Zone-C",
            "expected": "4-7 m",
            "expected_docs": ["3A", "SUP-B"],
            "scan_docs": ["3A", "SUP-B", "3B1"],
            "anchor_groups": [[r"Schreckstoff", r"警報物質"], [r"Zone[- ]?C"]],
            "checks": [("4-7 m", r"4\s*[-~]\s*7\s*m")],
        },
        {
            "name": "Mid-Strolling速度上限",
            "expected": "0.5-1.0 m/s",
            "expected_docs": ["3A"],
            "scan_docs": ["3A", "SUP-A"],
            "anchor_groups": [[r"Mid-Strolling"]],
            "checks": [("0.5-1.0 m/s", r"0\.5\s*[-~]\s*1\.0\s*m\s*/\s*s")],
        },
        {
            "name": "Zone-B 22°C超前量",
            "expected": "12-18天",
            "expected_docs": ["0D1", "1A"],
            "scan_docs": ["0D1", "1A", "4A", "4B"],
            "anchor_groups": [[r"B0-21", r"Zone-B"], [r"12\s*[-~]\s*18\s*天"]],
            "checks": [("12-18天", r"12\s*[-~]\s*18\s*天")],
        },
        {
            "name": "Zone-B Eh<0mV首觸",
            "expected": "5月下旬",
            "expected_docs": ["0D1", "4A", "4B"],
            "scan_docs": ["0D1", "4A", "4B"],
            "anchor_groups": [[r"B0-22", r"Zone-B"], [r"Eh\s*<\s*0\s*mV", r"5\s*月\s*下\s*旬"]],
            "checks": [("5月下旬", r"5\s*月\s*下\s*旬")],
        },
        {
            "name": "Fe2+安全距離",
            "expected": "40 cm",
            "expected_docs": ["1B1", "SUP-B"],
            "scan_docs": ["1B1", "SUP-B"],
            "anchor_groups": [[r"Fe2\+", r"Fe2\+", r"亞鐵", r"Fe²\+"], [r"40\s*cm"]],
            "checks": [("40 cm", r"40\s*cm")],
        },
        {
            "name": "H2S安全距離 靜水",
            "expected": "55-65 cm",
            "expected_docs": ["1B1", "SUP-B"],
            "scan_docs": ["1B1", "SUP-B"],
            "anchor_groups": [[r"H2S", r"硫化氫"], [r"靜水"], [r"55\s*[-~]\s*65\s*cm"]],
            "checks": [("55-65 cm", r"55\s*[-~]\s*65\s*cm")],
        },
        {
            "name": "H2S安全距離 微流/基準",
            "expected": ">=86 cm",
            "expected_docs": ["1B1", "SUP-B"],
            "scan_docs": ["1B1", "SUP-B", "3B1"],
            "anchor_groups": [[r"H2S", r"硫化氫"], [r"微流", r"基準", r"baseline"], [r"86\s*cm", r">=\s*86\s*cm", r"至少\s*86\s*cm"]],
            "checks": [(">=86 cm", r">?=?\s*86\s*cm|至少\s*86\s*cm")],
        },
        {
            "name": "H2S安全距離 擾動/水車",
            "expected": "120-160 cm",
            "expected_docs": ["SUP-B", "3B1"],
            "scan_docs": ["SUP-B", "3B1"],
            "anchor_groups": [[r"H2S", r"硫化氫"], [r"擾動", r"水車", r"重啟"], [r"120\s*[-~]\s*160\s*cm"]],
            "checks": [("120-160 cm", r"120\s*[-~]\s*160\s*cm")],
        },
        {
            "name": "H2S廢棄巢穴死亡區",
            "expected": "r=1.6-2.3 m",
            "expected_docs": ["4A", "SUP-B", "4B"],
            "scan_docs": ["4A", "SUP-B", "4B"],
            "anchor_groups": [[r"H2S", r"硫化氫"], [r"廢棄巢穴", r"死亡區"], [r"1\.6\s*[-~]\s*2\.3\s*m"]],
            "checks": [("r=1.6-2.3 m", r"r\s*=\s*1\.6\s*[-~]\s*2\.3\s*m|1\.6\s*[-~]\s*2\.3\s*m")],
        },
        {
            "name": "DO臨界下限",
            "expected": "3-4 mg/L",
            "expected_docs": ["0D1"],
            "scan_docs": ["0D1", "1A", "4A"],
            "anchor_groups": [[r"B0-08", r"DO", r"溶氧"], [r"3\s*[-~]\s*4\s*mg\s*/\s*L"]],
            "checks": [("3-4 mg/L", r"3\s*[-~]\s*4\s*mg\s*/\s*L")],
        },
    ]

    phase1 = []
    for spec in param_specs:
        rows = []
        for doc_key, info in docs.items():
            result = classify_parameter(doc_key, info, spec)
            if result:
                rows.append((doc_key, result))
        phase1.append((spec, rows))

    block_specs = {
        "0A1": [("Findings", r"\bFindings\b"), ("Carry_Forward", r"Carry_Forward")],
        "0B": [("Findings", r"\bFindings\b"), ("Carry_Forward", r"Carry_Forward")],
        "0C1": [("Findings", r"\bFindings\b"), ("Carry_Forward", r"Carry_Forward")],
        "0D1": [("Baseline_Facts/B0-XX", r"Baseline_Facts|B0-\d{2}"), ("Waterbody_Model_Table", r"Waterbody_Model_Table"), ("Carry_Forward", r"Carry_Forward")],
        "1A": [("Inherited_Baseline", r"Inherited_Baseline"), ("V1X-XX Findings", r"V1A-\d{2}"), ("Carry_Forward_To", r"Carry_Forward_To")],
        "1B1": [("Inherited_Baseline", r"Inherited_Baseline"), ("V1X-XX Findings", r"V1B-\d{2}"), ("Carry_Forward_To", r"Carry_Forward_To")],
        "2A": [("Inherited_Baseline", r"Inherited_Baseline"), ("V2X-XX Findings", r"V2A-\d{2}"), ("Carry_Forward_To", r"Carry_Forward_To")],
        "2B": [("Inherited_Baseline", r"Inherited_Baseline"), ("V2X-XX Findings", r"V2B-\d{2}"), ("Carry_Forward_To", r"Carry_Forward_To")],
        "2C": [("Inherited_Baseline", r"Inherited_Baseline"), ("V2X-XX Findings", r"V2C-\d{2}"), ("Carry_Forward_To", r"Carry_Forward_To")],
        "3A": [("Inherited_Baseline", r"Inherited_Baseline"), ("V3A-XX", r"V3A-\d{2}"), ("CF-XX", r"CF-\d{2}"), ("Open_Assumptions", r"Open_Assumptions")],
        "3B1": [("Inherited_Baseline", r"Inherited_Baseline"), ("V3B-XX", r"V3B-\d{2}"), ("Unresolved_Dependencies", r"Unresolved_Dependencies")],
        "4A": [("Inherited_Baseline", r"Inherited_Baseline"), ("V4X-XX Findings", r"V4A-\d{2}"), ("Carry_Forward", r"Carry_Forward")],
        "4B": [("Inherited_Baseline", r"Inherited_Baseline"), ("V4X-XX Findings", r"V4B-\d{2}"), ("Carry_Forward", r"Carry_Forward")],
        "SUP-A": [("Metadata", r"\bMetadata\b"), ("Inherited_Baseline", r"Inherited_Baseline"), ("VSUP-XX", r"VSUP-A-?\d{2}"), ("Correction_Instructions", r"Correction_Instructions"), ("Carry_Forward", r"Carry_Forward"), ("Unresolved", r"Unresolved")],
        "SUP-B": [("Metadata", r"\bMetadata\b"), ("Inherited_Baseline", r"Inherited_Baseline"), ("VSUP-XX", r"VSUP-B-?\d{2}"), ("Correction_Instructions", r"Correction_Instructions"), ("Carry_Forward", r"Carry_Forward"), ("Unresolved", r"Unresolved")],
    }
    phase2 = OrderedDict((doc_key, [(label, "存在" if re_has(docs[doc_key]["text"], pat) else "缺失") for label, pat in checks]) for doc_key, checks in block_specs.items())

    phase3 = []
    supb_lines = docs["SUP-B"]["lines"]
    supb_correction = section_slice(supb_lines, ["Correction_Instructions"], ["Carry_Forward", "Unresolved"])
    supb_codes = normalize_code_set(CODE_RE.findall("\n".join(supb_correction)))
    phase3.append(("SUP-B", "Correction_Instructions 目標碼", ", ".join(supb_codes) if supb_codes else "未抽出任何目標碼"))
    for code in supb_codes:
        target = code_target(code)
        if not target:
            phase3.append(("SUP-B", code, "無法映射目標文件"))
            continue
        status = "已見於目標文件" if code_exists_in_doc(code, docs[target]["text"]) else "目標文件缺碼"
        detail = f"{target} -> {status}"
        supb_snips = snippets_for_code(code, supb_lines)
        target_snips = snippets_for_code(code, docs[target]["lines"])
        if supb_snips:
            detail += f" | SUP-B: {supb_snips[0]}"
        if target_snips:
            detail += f" | {target}: {target_snips[0]}"
        phase3.append(("SUP-B", code, detail))

    inherited_3b1 = section_slice(docs["3B1"]["lines"], ["Inherited_Baseline"], ["V3B-01", "Findings", "Unresolved_Dependencies"])
    inherited_3b1_codes = [c for c in normalize_code_set(CODE_RE.findall("\n".join(inherited_3b1))) if c.startswith(("B0-", "V", "VSUP")) and not c.startswith("V3B-")]
    phase3.append(("3B1", "Inherited_Baseline 引用碼", ", ".join(inherited_3b1_codes) if inherited_3b1_codes else "未抽出任何引用碼"))
    for code in inherited_3b1_codes:
        target = code_target(code)
        if not target:
            phase3.append(("3B1", code, "無法映射上游文件"))
            continue
        detail = f"{target} -> {'存在' if code_exists_in_doc(code, docs[target]['text']) else '缺失'}"
        target_snips = snippets_for_code(code, docs[target]["lines"])
        if target_snips:
            detail += f" | {target}: {target_snips[0]}"
        phase3.append(("3B1", code, detail))

    inherited_4b = section_slice(docs["4B"]["lines"], ["Inherited_Baseline"], ["V4B-01", "Findings", "Carry_Forward"])
    has_vsup_b08 = any(canonical_code(c) == "VSUP-B08" for c in CODE_RE.findall("\n".join(inherited_4b))) or has_any("\n".join(inherited_4b), ["VSUP-B08", "VSUP-B-08"])
    phase3.append(("4B", "VSUP-B08", "存在於 Inherited_Baseline" if has_vsup_b08 else "未見於 Inherited_Baseline"))

    phase4 = []
    scope_rules = {
        "0-series": (["0A1", "0B", "0C1", "0D1"], ["索餌", "攻擊", "皮質醇", "神經丘", "咬餌", "逃跑", "覓食", "產卵", "護巢"]),
        "0A/0B-extra": (["0A1", "0B"], ["棲位"]),
        "1-3": (["1A", "1B1", "2A", "2B", "2C", "3A", "3B1"], ["產卵", "孵化", "護巢", "築巢"]),
        "4-series": (["4A", "4B"], ["CFF", "表淺神經丘", "SNs", "Hz"]),
    }
    for rule_name, (doc_list, terms) in scope_rules.items():
        for doc_key in doc_list:
            for term in terms:
                snips = find_plain_snippets(docs[doc_key]["lines"], [term], radius=1, limit=2)
                if snips:
                    phase4.append((doc_key, rule_name, term, snips[0]))

    inherited_1a = section_slice(docs["1A"]["lines"], ["Inherited_Baseline"], ["V1A-01", "Findings", "Carry_Forward_To"])
    phase5 = OrderedDict([
        ("4B是否繼承H2S廢棄巢穴死亡區", {"status": "有" if has_any("\n".join(inherited_4b), ["VSUP-B08", "VSUP-B-08", "1.6-2.3 m", "廢棄巢穴", "死亡區"]) else "無", "evidence": find_plain_snippets(inherited_4b, ["VSUP-B08", "VSUP-B-08", "1.6-2.3", "廢棄巢穴", "死亡區"], radius=0, limit=3)}),
        ("4B是否反映B0-22提早Eh<0mV", {"status": "有" if has_any("\n".join(inherited_4b), ["B0-22", "Eh<0", "5月下旬"]) else "無", "evidence": find_plain_snippets(inherited_4b, ["B0-22", "Eh<0", "5月下旬"], radius=0, limit=3)}),
        ("1A是否有B0-21 Zone-B溫度超前", {"status": "有" if has_any("\n".join(inherited_1a), ["B0-21", "12-18天", "Zone-B"]) else "無", "evidence": find_plain_snippets(inherited_1a, ["B0-21", "12-18", "Zone-B"], radius=0, limit=3)}),
        ("2C是否有Zone-B相關描述", {"status": "有" if has_any(docs["2C"]["text"], ["Zone-B", "zone b", "B0-21", "B0-22"]) else "無", "evidence": find_plain_snippets(docs["2C"]["lines"], ["Zone-B", "B0-21", "B0-22"], radius=1, limit=5)}),
        ("3B1是否有VSUP-B12水車重啟風險段落", {"status": "有" if has_any(docs["3B1"]["text"], ["VSUP-B12", "重啟", "restart", "水車"]) and has_any(docs["3B1"]["text"], ["H2S", "H₂S"]) else "無", "evidence": find_plain_snippets(docs["3B1"]["lines"], ["VSUP-B12", "重啟", "restart", "水車", "H2S", "H₂S"], radius=1, limit=6)})
    ])

    high, medium, low, research = [], [], [], []
    for spec, rows in phase1:
        for doc_key, result in rows:
            if result["status"] in {"偏差", "缺失"} and doc_key in spec.get("expected_docs", []):
                high.append(f"{doc_key} {spec['name']} -> {result['status']} ({result['value']})；建議對齊 {spec['expected']}")
    for doc_key, checks in phase2.items():
        for label, status in checks:
            if status == "缺失":
                medium.append(f"{doc_key} 缺少輸出區塊 {label}")
    for src, code, detail in phase3:
        if "缺" in detail or "未見" in detail or "未抽出" in detail:
            (high if src == "3B1" else medium).append(f"{src} {code} -> {detail}")
    for doc_key, rule_name, term, snippet in phase4:
        medium.append(f"{doc_key} scope疑慮[{rule_name}] {term} -> {snippet}")
    if phase5["4B是否繼承H2S廢棄巢穴死亡區"]["status"] == "無":
        high.append("4B 未在 Inherited_Baseline 明確繼承 VSUP-B08 / H2S廢棄巢穴死亡區 r=1.6-2.3 m")
    if phase5["4B是否反映B0-22提早Eh<0mV"]["status"] == "無":
        high.append("4B 未在 Inherited_Baseline 明確帶入 B0-22 (Zone-B Eh<0mV 首觸 5月下旬)")
    if phase5["1A是否有B0-21 Zone-B溫度超前"]["status"] == "無":
        medium.append("1A 未明確帶入 B0-21 (Zone-B 22°C 超前 12-18天)")
    if phase5["2C是否有Zone-B相關描述"]["status"] == "無":
        low.append("2C 未見 Zone-B / B0-21 / B0-22 相關描述")
    if phase5["3B1是否有VSUP-B12水車重啟風險段落"]["status"] == "無":
        high.append("3B1 未見 VSUP-B12 / 水車重啟 + H2S 風險的明確段落")

    def dedupe(items):
        out, seen = [], set()
        for item in items:
            if item not in seen:
                seen.add(item)
                out.append(item)
        return out
    high, medium, low, research = map(dedupe, (high, medium, low, research))

    out = []
    out.append("## Phase 1 量化參數比對矩陣")
    out.append("")
    for spec, rows in phase1:
        out.append(f"### {spec['name']}")
        out.append(f"- 期望值：{spec['expected']}")
        out.append("- 出現文件：")
        for doc_key, result in rows:
            evidence = f" | 證據：{result['snippets'][0]}" if result['snippets'] else ""
            out.append(f"  - {doc_key}：{result['value']} — {result['status']}{evidence}")
        if not rows:
            out.append("  - 未找到任何文件")
        out.append("")

    out.append("## Phase 2 輸出區塊完整性")
    out.append("")
    for doc_key, checks in phase2.items():
        out.append(f"### {doc_key}")
        for label, status in checks:
            out.append(f"- {label}：{status}")
        out.append("")

    out.append("## Phase 3 引用鏈問題")
    out.append("")
    for src, code, detail in phase3:
        out.append(f"- {src} → {code} → {detail}")
    out.append("")

    out.append("## Phase 4 Scope 合規問題")
    out.append("")
    if phase4:
        for doc_key, rule_name, term, snippet in phase4:
            out.append(f"- {doc_key} [{rule_name}] {term}：{snippet}")
    else:
        out.append("- 未檢出明顯 scope 關鍵字衝突")
    out.append("")

    out.append("## Phase 5 特別重點檢查")
    out.append("")
    for title, item in phase5.items():
        evidence = f"；證據：{item['evidence'][0]}" if item['evidence'] else ""
        out.append(f"- {title}：{item['status']}{evidence}")
    out.append("")

    out.append("## 彙整問題清單（按優先級）")
    out.append("")
    out.append("### 高優先（矛盾/缺失影響戰術建議）")
    out.extend([f"- {x}" for x in high] or ["- 無"])
    out.append("")
    out.append("### 中優先（格式缺漏/引用不完整）")
    out.extend([f"- {x}" for x in medium] or ["- 無"])
    out.append("")
    out.append("### 低優先（小細節/僅供記錄）")
    out.extend([f"- {x}" for x in low] or ["- 無"])
    out.append("")
    out.append("### 需要 Deep Research（非 Python patch 可解）")
    out.extend([f"- {x}" for x in research] or ["- 無"])
    print("\n".join(out))


if __name__ == "__main__":
    main()
